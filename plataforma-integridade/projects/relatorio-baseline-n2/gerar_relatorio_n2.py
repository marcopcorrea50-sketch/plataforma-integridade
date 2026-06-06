"""
GERADOR DE RELATÓRIO BASELINE N2 - PADRÃO VALE
Adaptado para Transportador de Longa Distância (TCLD) e equipamentos similares.

Formato baseado no modelo RL_RC2092KS_06_N2.docx:
- Cabeçalho corporativo OPERAÇÃO – MANUTENÇÃO.
- 6 seções: Objetivo / Referências / Metodologia / Resumo / Conclusão / Recomendações.
- Componentes estruturais configuráveis (estilo Caption em maiúsculas).
- Layout de fotos configurável por componente (1, 2, 4 ou 6 fotos/página).
- Recomendações gerais + recomendações vinculadas a figuras específicas.
- Importação automática de fotos por nome de arquivo.

Uso:
    python3 gerar_relatorio_n2.py --config config.json --fotos ./fotos --saida relatorio.docx
    python3 gerar_relatorio_n2.py --config config.json --fotos ./fotos --saida relatorio.docx --pdf
"""

from __future__ import annotations
import argparse, csv, json, os, re, subprocess, sys
from dataclasses import dataclass, field
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRO: pip install python-docx Pillow", file=sys.stderr)
    sys.exit(1)


# =============================================================
# TEMPLATES DE COMPONENTES (TLD = transportador, RC = recuperadora)
# =============================================================

COMPONENTES_TLD = [
    "INSPEÇÃO ESTRUTURAL DA CABEÇA DE ACIONAMENTO",
    "ESTRUTURA METÁLICA - CAVALETES",
    "ESTRUTURA METÁLICA - LONGARINAS E TRANSVERSINAS",
    "SUPORTES DE ROLETES E ESTRUTURAS AUXILIARES",
    "GALERIA DE PROTEÇÃO",
    "PASSARELAS, ESCADAS E PLATAFORMAS",
    "GUARDA-CORPOS E RODAPÉS",
    "CURVAS VERTICAIS E HORIZONTAIS",
    "CONJUNTO DE ESTICAMENTO",
    "CONJUNTO DE CAUDA / RETORNO",
    "ESTAÇÕES DE TRANSFERÊNCIA E CHUTES",
    "FUNDAÇÕES E CHUMBADORES",
    "SISTEMA DE PROTEÇÃO ANTICORROSIVA",
    "JUNTAS PARAFUSADAS",
    "REGIÕES SOLDADAS",
]

COMPONENTES_RC = [
    "INSPEÇÃO ESTRUTURAL DO SISTEMA DE TRANSLAÇÃO",
    "SISTEMA DE TRANSLAÇÃO LADO ESQUERDO",
    "SISTEMA DE TRANSLAÇÃO LADO DIREITO",
    "VIGA EQUALIZADORA / BALANCIM DE RODAS",
    "ESTRUTURA DE ACESSO À MESA DE GIRO DA LANÇA",
    "ESTRUTURA DA MESA DE GIRO DA LANÇA",
    "ESTRUTURA DAS RODAS EQUALIZADORAS",
    "ESTRUTURA DO PORTAL",
    "ESTRUTURA DA CONTRA LANÇA",
    "ESTRUTURA DO CARRO TENSOR DA LANÇA",
    "INSPEÇÕES NAS REGIÕES ARTICULADAS DO MASTRO E TALAS",
    "INSPEÇÃO DAS ESTRUTURAS DA CABINE",
    "INSPEÇÃO ESTRUTURAL DA LANÇA",
    "GALERIA DA LANÇA - LADO ESQUERDO",
    "GALERIA DA LANÇA - LADO DIREITO",
    "INSPEÇÃO ESTRUTURAL DOS ELEMENTOS INFERIORES DA LANÇA",
]


# =============================================================
# DATA CLASSES
# =============================================================

@dataclass
class Foto:
    arquivo: str
    legenda: str = ""
    ordem: int = 0
    recomendacao: str = ""  # recomendação específica desta figura (opcional)
    fig_num: int = 0  # preenchido durante a geração


@dataclass
class Componente:
    titulo: str  # ex: "INSPEÇÃO ESTRUTURAL DA CABEÇA DE ACIONAMENTO"
    introducao: str = ""
    fotos: list[Foto] = field(default_factory=list)
    layout: str = "1"  # '1', '2', '4', '6'


@dataclass
class Config:
    # Cabeçalho corporativo
    cabec_central: str = "INTEGRIDADE ESTRUTURAL"
    classificacao: str = "USO INTERNO"
    projeto_top: str = "PROJETO FERRO S11D"
    tipo_relatorio: str = "OPERACAO – MANUTENCAO"   # Linha 1 esquerda
    patio: str = ""                                  # Linha 2: ex "5° PATIO DE PRODUTO AREA 2092"
    circuito: str = ""                               # Linha 3: ex "CIRCUITO 2092KS/TCLD 01"
    descricao_inspecao: str = ""                     # Linha 4: ex "INSPECAO VISUAL – TLD-001"
    relatorio_subtipo: str = "RELATÓRIO TÉCNICO – N2"
    n_vale: str = ""
    n_doc: str = ""
    revisao: str = "A"
    codigo_template: str = "PE-G-606 - Rev. 5"
    # Conteúdo
    objetivo: str = ""
    normas: list[str] = field(default_factory=list)
    metodologia: str = ""
    materiais_equipamentos: list[str] = field(default_factory=list)
    desenhos: list[str] = field(default_factory=list)
    nota_descritiva: str = ""                        # texto descritivo do ativo
    fotos_panoramicas: list[Foto] = field(default_factory=list)
    componentes: list[Componente] = field(default_factory=list)
    resumo: str = ""
    conclusao: str = ""
    recomendacoes: list[str] = field(default_factory=list)
    revisoes: list[dict] = field(default_factory=list)
    # Campos opcionais usados pelos templates Vale
    planos: list = field(default_factory=list)
    planos_texto_contexto: str = ""


# =============================================================
# HELPERS XML
# =============================================================

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_table_borders(table, color="000000", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz); b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_field(paragraph, instr):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(i); run._r.append(f2)

def add_seq(paragraph, name):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve')
    i.text = f' SEQ {name} \\* ARABIC '
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(i); run._r.append(f2)

def add_toc(paragraph, switch):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve')
    i.text = f' TOC {switch} '
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(i); run._r.append(f2)


# =============================================================
# GERADOR
# =============================================================

class N2ReportBuilder:
    def __init__(self, cfg: Config, base_fotos: Path):
        self.cfg = cfg
        self.base_fotos = base_fotos
        self.doc = Document()
        self._fig_count = 0
        self._fig_map = {}  # arquivo -> n da figura (para recomendações vinculadas)
        self._setup_pagina()
        self._setup_estilos()
        self._add_cabecalho()
        self._add_rodape()

    def _setup_pagina(self):
        for s in self.doc.sections:
            s.page_width = Cm(21.0); s.page_height = Cm(29.7)
            s.top_margin = Cm(4.8); s.bottom_margin = Cm(1.8)
            s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)
            s.header_distance = Cm(0.5); s.footer_distance = Cm(0.5)

    def _setup_estilos(self):
        st = self.doc.styles['Normal']
        st.font.name = 'Arial'; st.font.size = Pt(11)
        pf = st.paragraph_format
        pf.line_spacing = 1.15; pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.0)
        for name, sz in [("Heading 1", 14), ("Heading 2", 12)]:
            h = self.doc.styles[name]
            h.font.name = 'Arial'; h.font.size = Pt(sz); h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(18); h.paragraph_format.space_after = Pt(12)
            h.paragraph_format.first_line_indent = Cm(0)
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -------------------------------------------------------
    # CABECALHO (formato OPERAÇÃO-MANUTENÇÃO N2)
    # -------------------------------------------------------
    def _add_cabecalho(self):
        header = self.doc.sections[0].header
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)
        # =====================================================
        # TABELA 1 (topo): 1 linha x 3 colunas
        # [logo] | INTEGRIDADE ESTRUTURAL + USO INTERNO | PROJETO FERRO S11D
        # =====================================================
        t1 = header.add_table(rows=1, cols=3, width=Cm(17))
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t1, "000000", "6")
        widths1 = [Cm(2.0), Cm(10.5), Cm(4.5)]
        for i, w in enumerate(widths1):
            t1.rows[0].cells[i].width = w
        # c0 vazio
        c = t1.rows[0].cells[0]; c.text = ""
        # c1 central
        c = t1.rows[0].cells[1]; c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(self.cfg.cabec_central); r.bold = True; r.font.size=Pt(11); r.font.name='Arial'
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.first_line_indent = Cm(0)
        r = p2.add_run(self.cfg.classificacao); r.bold = True; r.font.size=Pt(11); r.font.name='Arial'
        # c2 PROJETO FERRO S11D
        c = t1.rows[0].cells[2]; c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(self.cfg.projeto_top); r.bold = True; r.font.size=Pt(11); r.font.name='Arial'
        for cell in t1.rows[0].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # =====================================================
        # TABELA 2 (descrição): 2 linhas x 3 colunas
        # descrição (rowspan 2) | NºVALE+Nº | PÁGINA+REV
        # =====================================================
        t = header.add_table(rows=2, cols=3, width=Cm(17))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t, "000000", "6")
        cw = [Cm(10.5), Cm(2.5), Cm(4.0)]
        for r_idx in range(2):
            for i, w in enumerate(cw):
                t.rows[r_idx].cells[i].width = w

        # Coluna 0 mesclada verticalmente
        desc_cell = t.rows[0].cells[0].merge(t.rows[1].cells[0])
        desc_cell.text = ""
        linhas_desc = [
            self.cfg.tipo_relatorio,
            self.cfg.patio,
            self.cfg.circuito,
            self.cfg.descricao_inspecao,
            self.cfg.relatorio_subtipo,
        ]
        for idx, txt in enumerate(linhas_desc):
            p = desc_cell.paragraphs[0] if idx == 0 else desc_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(txt or " "); r.bold = True; r.font.name='Arial'; r.font.size=Pt(9)

        # Coluna 1 (linha 0): Nº VALE + valor
        c = t.rows[0].cells[1]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Nº VALE"); r.font.name='Arial'; r.font.size=Pt(9)
        if self.cfg.n_vale:
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.first_line_indent = Cm(0)
            r = p2.add_run(self.cfg.n_vale); r.bold=True; r.font.name='Arial'; r.font.size=Pt(9)

        # Coluna 2 (linha 0): PÁGINA + numeração
        c = t.rows[0].cells[2]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run("PÁGINA"); r.font.name='Arial'; r.font.size=Pt(9)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Cm(0)
        add_field(p2, "PAGE \\* MERGEFORMAT")
        p2.add_run("/")
        add_field(p2, "NUMPAGES \\* MERGEFORMAT")
        for run in p2.runs:
            run.bold = True; run.font.name='Arial'; run.font.size=Pt(9)

        # Coluna 1 (linha 1): Nº + valor
        c = t.rows[1].cells[1]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Nº "); r.font.name='Arial'; r.font.size=Pt(9)
        if self.cfg.n_doc:
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.first_line_indent = Cm(0)
            r = p2.add_run(self.cfg.n_doc); r.bold=True; r.font.name='Arial'; r.font.size=Pt(9)

        # Coluna 2 (linha 1): REV + valor
        c = t.rows[1].cells[2]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run("REV."); r.font.name='Arial'; r.font.size=Pt(9)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Cm(0)
        r = p2.add_run(self.cfg.revisao); r.bold = True; r.font.name='Arial'; r.font.size=Pt(10)

        for row in t.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_rodape(self):
        footer = self.doc.sections[0].footer
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        t = footer.add_table(rows=1, cols=3, width=Cm(17))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = [Cm(5), Cm(7), Cm(5)]
        for i, w in enumerate(widths):
            t.rows[0].cells[i].width = w
            remove_cell_borders(t.rows[0].cells[i])
        # Esquerda
        c = t.rows[0].cells[0]; c.text = ""
        p = c.paragraphs[0]; p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(self.cfg.codigo_template)
        r.font.name='Arial'; r.font.size=Pt(8); r.italic=True
        # Centro: número
        c = t.rows[0].cells[1]; c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
        add_field(p, "PAGE \\* MERGEFORMAT")
        for run in p.runs:
            run.font.name='Arial'; run.font.size=Pt(9)

    # -------------------------------------------------------
    # HELPERS DE CONTEUDO
    # -------------------------------------------------------
    def H1(self, txt, num=None):
        p = self.doc.add_paragraph(style='Heading 1')
        text = f"{num}. {txt}" if num is not None else txt
        r = p.add_run(text)
        r.bold = True; r.font.name='Arial'; r.font.size=Pt(14)
        return p

    def H2(self, txt, num=None):
        p = self.doc.add_paragraph(style='Heading 2')
        text = f"{num} {txt}" if num is not None else txt
        r = p.add_run(text)
        r.bold = True; r.font.name='Arial'; r.font.size=Pt(12)
        return p

    def H3(self, txt, num=None):
        p = self.doc.add_paragraph(style='Heading 3')
        text = f"{num} {txt}" if num is not None else txt
        r = p.add_run(text)
        r.bold = True; r.font.name='Arial'; r.font.size=Pt(11)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
        return p

    def P(self, txt="", bold=False, italic=False, align=None, indent=True, size=11):
        p = self.doc.add_paragraph()
        if not indent:
            p.paragraph_format.first_line_indent = Cm(0)
        if align:
            p.alignment = align
        r = p.add_run(txt)
        r.bold = bold; r.italic = italic
        r.font.name='Arial'; r.font.size=Pt(size)
        return p

    def BULLET(self, txt):
        p = self.doc.add_paragraph(txt, style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            run.font.name='Arial'; run.font.size=Pt(11)
        return p

    def NUM(self, txt):
        p = self.doc.add_paragraph(txt, style='List Number')
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            run.font.name='Arial'; run.font.size=Pt(11)
        return p

    def PB(self):
        self.doc.add_page_break()

    def CAPTION_SUBSECAO(self, texto):
        """Caption em maiúsculas (estilo dos componentes do N2)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        try:
            p.style = self.doc.styles['Caption']
        except KeyError:
            pass
        r = p.add_run(texto.upper())
        r.bold = True; r.font.name='Arial'; r.font.size=Pt(12)
        return p

    # -------------------------------------------------------
    # FOTOS / PRANCHAS / LEGENDAS
    # -------------------------------------------------------
    def _add_legenda_figura(self, foto: Foto, sufixo_em_itens=False):
        self._fig_count += 1
        foto.fig_num = self._fig_count
        if foto.arquivo:
            self._fig_map[foto.arquivo] = self._fig_count
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        try:
            p.style = self.doc.styles['Caption']
        except KeyError:
            pass
        r = p.add_run("Figura ")
        r.font.name='Arial'; r.font.size=Pt(10)
        add_seq(p, "Figura")
        legenda = foto.legenda or "Registro fotográfico."
        r = p.add_run(f" – {legenda}")
        r.font.name='Arial'; r.font.size=Pt(10)

    def _inserir_foto(self, foto: Foto, largura_cm: float = 16, altura_cm: float | None = None):
        path = self._resolve_foto(foto.arquivo)
        if not path or not os.path.exists(path):
            self.P(f"[Foto não encontrada: {foto.arquivo}]", italic=True, indent=False)
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        try:
            if altura_cm:
                p.add_run().add_picture(path, width=Cm(largura_cm), height=Cm(altura_cm))
            else:
                p.add_run().add_picture(path, width=Cm(largura_cm))
        except Exception as e:
            p.add_run(f"[Erro: {e}]")

    def _add_prancha_componente(self, fotos: list[Foto], layout: str):
        """Adiciona fotos do componente conforme layout escolhido.
        Layout '1': 1 foto + 1 legenda por página.
        Layouts '2','4','6': prancha com legenda única.
        """
        if not fotos:
            return
        if layout == "1":
            for foto in fotos:
                self._inserir_foto(foto, largura_cm=16)
                self._add_legenda_figura(foto)
            return
        # Prancha
        config = {"2": (1, 2, 7.5, 10), "4": (2, 2, 7.5, 9), "6": (3, 2, 7.5, 7)}
        rows, cols, w_cm, h_cm = config.get(layout, config["6"])
        per_page = rows * cols
        for start in range(0, len(fotos), per_page):
            chunk = fotos[start:start + per_page]
            t = self.doc.add_table(rows=rows, cols=cols)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl = t._tbl; tblPr = tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), 'B0B0B0')
                tblBorders.append(b)
            tblPr.append(tblBorders)
            for idx, foto in enumerate(chunk):
                r = idx // cols; c = idx % cols
                cell = t.rows[r].cells[c]; cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                path = self._resolve_foto(foto.arquivo)
                if path and os.path.exists(path):
                    try:
                        p.add_run().add_picture(path, width=Cm(w_cm), height=Cm(h_cm))
                    except Exception as e:
                        p.add_run(f"[erro: {e}]")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # legenda única (usa a primeira foto do chunk como referência)
            self._add_legenda_figura(chunk[0])

    def _resolve_foto(self, arquivo):
        if not arquivo:
            return ""
        if os.path.isabs(arquivo) and os.path.exists(arquivo):
            return arquivo
        cand = self.base_fotos / arquivo
        if cand.exists():
            return str(cand)
        for p in self.base_fotos.rglob(arquivo):
            return str(p)
        return arquivo

    # -------------------------------------------------------
    # CONSTRUÇÃO DO CORPO - estrutura ABNT em 8 seções
    # -------------------------------------------------------
    def build(self):
        self._tabela_revisoes()
        self.PB()
        self._indice()
        self.PB()
        self._indice_figuras()
        self.PB()
        # 1. OBJETIVO
        self._sec_1_objetivo()
        # 2. REFERÊNCIAS E NORMAS
        self._sec_2_referencias()
        # 3. METODOLOGIA
        self._sec_3_metodologia()
        # 4. PLANOS
        self._sec_4_planos()
        # 5. DESENHOS DE REFERÊNCIA
        self._sec_5_desenhos()
        # 6. CIRCUITOS E EVIDÊNCIAS
        self._sec_6_circuitos_evidencias()
        # 7. CONCLUSÃO
        self._sec_7_conclusao()
        # 8. RECOMENDAÇÕES
        self._sec_8_recomendacoes()

    def _tabela_revisoes(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("REVISÕES")
        r.bold = True; r.italic = True; r.font.name='Arial'; r.font.size=Pt(12)
        # tipos
        for txt in [
            ("TE: TIPO   ", True, "A - PRELIMINAR    C - PARA CONHECIMENTO    E - PARA CONSTRUÇÃO    G - CONFORME CONSTRUÍDO"),
            ("EMISSÃO   ", True, "B - PARA APROVAÇÃO    D - PARA COTAÇÃO    F - CONFORME COMPRADO    H - CANCELADO"),
        ]:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(txt[0]); r.bold = True; r.font.name='Arial'; r.font.size=Pt(8)
            r = p.add_run(txt[2]); r.font.name='Arial'; r.font.size=Pt(8)
        # tabela
        t = self.doc.add_table(rows=18, cols=8)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t, "000000", "6")
        headers = ["Rev.", "TE", "Descrição", "Por", "Ver.", "AP.", "Aut.", "Data"]
        ws = [Cm(1.3), Cm(1.0), Cm(7.0), Cm(1.5), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.8)]
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            for run in c.paragraphs[0].runs:
                run.bold = True; run.font.name='Arial'; run.font.size=Pt(10)
            c.width = ws[i]
        for ri, rev in enumerate(self.cfg.revisoes[:17], start=1):
            vals = [rev.get(k, "") for k in ("rev","te","descricao","por","ver","ap","aut","data")]
            for j, v in enumerate(vals):
                c = t.rows[ri].cells[j]; c.text = str(v)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 2 else WD_ALIGN_PARAGRAPH.LEFT
                c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                for run in c.paragraphs[0].runs:
                    run.font.name='Arial'; run.font.size=Pt(10)
                c.width = ws[j]
        for ri in range(1, 18):
            t.rows[ri].height = Cm(0.7)
            t.rows[ri].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    def _indice(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("ÍNDICE")
        r.bold = True; r.font.name='Arial'; r.font.size=Pt(14)
        self.P("")
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
        r = p.add_run("DESCRIÇÃO"); r.bold = True; r.font.name='Arial'; r.font.size=Pt(10)
        p.add_run("\t")
        r = p.add_run("PÁGINA"); r.bold = True; r.font.name='Arial'; r.font.size=Pt(10)
        self.P("")
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        add_toc(p, '\\o "1-3" \\h \\z \\u')
        p = self.doc.add_paragraph("(Atualize com F9 no Word)")
        p.paragraph_format.first_line_indent = Cm(0)
        for run in p.runs:
            run.italic = True; run.font.color.rgb = RGBColor.from_string("808080"); run.font.size = Pt(9)

    def _indice_figuras(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("ÍNDICE DE FIGURAS")
        r.bold = True; r.font.name='Arial'; r.font.size=Pt(14)
        self.P("")
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        add_toc(p, '\\h \\z \\c "Figura"')

    # =========================================================
    # SEÇÕES ABNT (estrutura 8 seções com hierarquia 1.1, 1.1.1)
    # =========================================================

    def _sec_1_objetivo(self):
        self.H1("1. OBJETIVO")
        self.H2("1.1 ESCOPO DA INSPEÇÃO")
        self.P(self.cfg.objetivo or
            "Este relatório visa identificar e documentar a condição inicial (baseline) das estruturas metálicas do ativo inspecionado, estabelecendo a condição de referência para acompanhamento futuro da integridade do equipamento e servindo como base para futuras comparações, manutenções e intervenções corretivas.")

    def _sec_2_referencias(self):
        self.H1("2. REFERÊNCIAS E NORMAS")
        self.H2("2.1 NORMAS E PROCEDIMENTOS APLICÁVEIS")
        normas = self.cfg.normas or [
            "ABNT NBR. 8800 – Norma Brasileira para Estruturas Metálicas.",
            "Norma de Soldagem AWS D1.1 Ed 2015.",
            "PRO-024969 – Inspeção Integridade Estrutural.",
            "PNR- 000048 – Estruturas de Aço.",
            "ABNT NBR 4628-3 - Avaliação do grau de enferrujamento.",
            "PGS-005273 – Proteção anticorrosiva de pintura.",
            "PNR-00047 – Integridade Estrutural.",
        ]
        for n in normas:
            self.BULLET(n)

    def _sec_3_metodologia(self):
        self.H1("3. METODOLOGIA")
        self.H2("3.1 CRITÉRIOS DE INSPEÇÃO VISUAL")
        self.P(self.cfg.metodologia or
            "A inspeção visual baseline das estruturas metálicas foi realizada seguindo procedimento detalhado: inspeção visual minuciosa por trecho identificando sinais de desgaste, corrosão, trincas, deformações e flambagem; verificação de juntas parafusadas e ligações soldadas; avaliação do sistema de proteção anticorrosiva conforme ISO 4628-3; documentação fotográfica de toda evidência relevante. A inspeção foi conduzida com EPIs adequados (NR-12, NR-35) e com emissão prévia de Permissão de Trabalho e Análise Preliminar de Risco.")
        self.H2("3.2 MATERIAIS E EQUIPAMENTOS UTILIZADOS")
        mats = self.cfg.materiais_equipamentos or ["Câmera fotográfica", "Lanterna", "Trena", "Registro de campo"]
        for m in mats:
            self.BULLET(m)

    def _sec_4_planos(self):
        self.H1("4. LISTA DOS PLANOS DE MANUTENÇÃO E LOCAL DE INSTALAÇÃO")
        self.H2("4.1 PLANOS E LOCAIS CADASTRADOS")
        if self.cfg.planos_texto_contexto:
            self.P(self.cfg.planos_texto_contexto)
        if self.cfg.planos:
            t = self.doc.add_table(rows=len(self.cfg.planos)+1, cols=4)
            set_table_borders(t, "000000", "6")
            headers = ["Centro trab. respons.", "Plano de manutenção", "Local de instalação", "Denominação do loc.instalação"]
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]; c.text = h
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                for run in c.paragraphs[0].runs:
                    run.bold = True; run.font.name='Arial'; run.font.size=Pt(9)
                set_cell_bg(c, "D9E1F2")
            for ri, plano in enumerate(self.cfg.planos, start=1):
                vals = [plano.centro, plano.plano, plano.local, plano.denominacao]
                for j, v in enumerate(vals):
                    c = t.rows[ri].cells[j]; c.text = v
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j < 2 else WD_ALIGN_PARAGRAPH.LEFT
                    c.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                    for run in c.paragraphs[0].runs:
                        run.font.name='Arial'; run.font.size=Pt(9)
        else:
            self.P("[Lista de planos a ser inserida]", italic=True)

    def _sec_5_desenhos(self):
        self.H1("5. DESENHOS DE REFERÊNCIA")
        # 5.1 Lista de desenhos
        self.H2("5.1 LISTA DE DESENHOS")
        if self.cfg.desenhos:
            cols = 2
            n = len(self.cfg.desenhos)
            rows = (n + cols - 1) // cols
            t = self.doc.add_table(rows=rows, cols=cols)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, d in enumerate(self.cfg.desenhos):
                r = i // cols; c = i % cols
                cell = t.rows[r].cells[c]; cell.text = d
                cell.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                for run in cell.paragraphs[0].runs:
                    run.font.name='Arial'; run.font.size=Pt(10)
        else:
            self.P("[Lista de desenhos a ser inserida]", italic=True)
        # 5.2 Registros dos desenhos (descrição + panorâmicas)
        self.H2("5.2 REGISTROS DOS DESENHOS")
        if self.cfg.nota_descritiva:
            self.P(self.cfg.nota_descritiva)
        for f in self.cfg.fotos_panoramicas:
            self._inserir_foto(f, largura_cm=15)
            self._add_legenda_figura(f)

    def _sec_6_circuitos_evidencias(self):
        self.H1("6. CIRCUITOS E EVIDÊNCIAS")
        # Cada componente vira uma subseção 6.1, 6.2, 6.3...
        for ci, comp in enumerate(self.cfg.componentes, start=1):
            sec_num = f"6.{ci}"
            self.H2(f"{sec_num} {comp.titulo}")
            if comp.introducao:
                self.P(comp.introducao)
            layout = comp.layout if comp.layout in ("1", "2", "4", "6") else "1"
            # Cada foto vira uma sub-sub-seção 6.X.Y
            if layout == "1":
                # 1 foto por página - cada uma com sua sub-sub-seção
                for fi, foto in enumerate(comp.fotos, start=1):
                    titulo_evid = self._titulo_evidencia(foto)
                    self.H3(f"{sec_num}.{fi} {titulo_evid}")
                    self._inserir_foto(foto, largura_cm=16)
                    self._add_legenda_figura(foto)
            else:
                # Pranchas - cada prancha como uma sub-sub-seção
                config = {"2": (1, 2, 7.5, 10), "4": (2, 2, 7.5, 9), "6": (3, 2, 7.5, 7)}
                rows, cols, w_cm, h_cm = config.get(layout, config["6"])
                per_page = rows * cols
                for pi, start in enumerate(range(0, len(comp.fotos), per_page), start=1):
                    chunk = comp.fotos[start:start + per_page]
                    titulo_prancha = self._titulo_evidencia(chunk[0]) if chunk else f"Prancha {pi}"
                    self.H3(f"{sec_num}.{pi} {titulo_prancha}")
                    self._add_prancha_chunk(chunk, rows, cols, w_cm, h_cm)
                    self._add_legenda_figura(chunk[0])

    def _titulo_evidencia(self, foto):
        """Gera título da sub-sub-seção a partir da legenda da foto."""
        if foto.legenda:
            # pega os primeiros 60 chars da legenda
            t = foto.legenda.strip()
            return t[:80] + ('...' if len(t) > 80 else '')
        if foto.arquivo:
            return foto.arquivo.replace('_', ' ').rsplit('.', 1)[0]
        return "Evidência"

    def _add_prancha_chunk(self, chunk, rows, cols, w_cm, h_cm):
        """Adiciona uma prancha individual sem legenda (legenda é adicionada separadamente)."""
        t = self.doc.add_table(rows=rows, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl = t._tbl; tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), 'B0B0B0')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        for idx, foto in enumerate(chunk):
            r = idx // cols; c = idx % cols
            cell = t.rows[r].cells[c]; cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            path = self._resolve_foto(foto.arquivo)
            if path and os.path.exists(path):
                try:
                    p.add_run().add_picture(path, width=Cm(w_cm), height=Cm(h_cm))
                except Exception as e:
                    p.add_run(f"[erro: {e}]")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _sec_7_conclusao(self):
        self.H1("7. CONCLUSÃO")
        self.H2("7.1 RESULTADO CONSOLIDADO")
        self.P(self.cfg.conclusao or
            "A condição geral do ativo é satisfatória nas seções inspecionadas, sem sinais visíveis de deformações ou desgaste significativo na maior parte das estruturas, indicando que o equipamento está em estado adequado para uso como referência baseline. Os pontos de atenção identificados, quando existentes, estão detalhados nas recomendações específicas por figura na próxima seção.")

    def _sec_8_recomendacoes(self):
        self.H1("8. RECOMENDAÇÕES")
        self.H2("8.1 AÇÕES RECOMENDADAS")
        # gerais
        if self.cfg.recomendacoes:
            self.P("Com base nas inspeções realizadas, recomenda-se:")
            for r in self.cfg.recomendacoes:
                self.NUM(r)
        # por figura
        recs_por_figura = []
        for comp in self.cfg.componentes:
            for foto in comp.fotos:
                if foto.recomendacao and foto.fig_num:
                    recs_por_figura.append((foto.fig_num, foto.arquivo, foto.recomendacao, foto.legenda))
        recs_por_figura.sort(key=lambda x: x[0])
        if recs_por_figura:
            self.P()
            self.P("Recomendações vinculadas a evidências fotográficas específicas:", bold=True, indent=False)
            for fig_num, arq, rec, leg in recs_por_figura:
                p = self.doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(f"Recomendação para Figura {fig_num}:")
                r.bold = True; r.font.name='Arial'; r.font.size=Pt(11)
                if leg:
                    p2 = self.doc.add_paragraph()
                    p2.paragraph_format.first_line_indent = Cm(0)
                    r = p2.add_run(f"({leg})")
                    r.italic = True; r.font.name='Arial'; r.font.size=Pt(10)
                    r.font.color.rgb = RGBColor.from_string("606060")
                self.P(rec)



# =============================================================
# IMPORTACAO DE FOTOS
# =============================================================

FNAME_RE = re.compile(r"^(?P<comp>[A-Z0-9\-]+)_(?P<sub>[A-Z0-9\-]+)_(?P<nn>\d+)\.(jpg|jpeg|png)$", re.IGNORECASE)

def importar_fotos_por_pasta(pasta: Path, componentes: list[Componente]):
    if not pasta.exists():
        return 0
    contador = 0
    for f in sorted(pasta.iterdir()):
        if f.suffix.lower() not in {'.jpg', '.jpeg', '.png'}:
            continue
        m = FNAME_RE.match(f.name)
        if not m:
            continue
        prefixo = m.group('comp').upper()
        for comp in componentes:
            chave = re.sub(r'[^A-Z0-9]', '', comp.titulo.upper())
            if prefixo in chave:
                if not any(ff.arquivo == f.name for ff in comp.fotos):
                    comp.fotos.append(Foto(arquivo=f.name, legenda=""))
                    contador += 1
                break
    return contador


def importar_fotos_csv(caminho: Path, componentes: list[Componente]):
    if not caminho.exists():
        return 0
    contador = 0
    with open(caminho, encoding='utf-8-sig', newline='') as fh:
        sample = fh.read(2048); fh.seek(0)
        sep = ';' if sample.count(';') > sample.count(',') else ','
        rd = csv.DictReader(fh, delimiter=sep)
        rows = list(rd)
    for row in rows:
        comp_titulo = (row.get('componente') or '').strip()
        comp = next((c for c in componentes if c.titulo == comp_titulo), None)
        if not comp:
            comp = Componente(titulo=comp_titulo, layout=row.get('layout', '1') or '1')
            componentes.append(comp)
        elif row.get('layout'):
            comp.layout = row.get('layout', comp.layout)
        comp.fotos.append(Foto(
            arquivo=(row.get('arquivo') or '').strip(),
            legenda=(row.get('legenda') or '').strip(),
            ordem=int(row.get('ordem') or 0),
            recomendacao=(row.get('recomendacao') or '').strip()
        ))
        contador += 1
    for c in componentes:
        c.fotos.sort(key=lambda f: f.ordem)
    return contador


def converter_pdf(docx_path):
    out_dir = os.path.dirname(os.path.abspath(docx_path))
    for cmd in (['soffice'], ['libreoffice']):
        try:
            subprocess.run(cmd + ['--headless', '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
                          check=True, capture_output=True, timeout=120)
            pdf_path = docx_path.replace('.docx', '.pdf')
            if os.path.exists(pdf_path):
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired, subprocess.CalledProcessError):
            continue
    return None


def main():
    ap = argparse.ArgumentParser(description="Gerador N2 ABNT - Padrao Vale.")
    ap.add_argument('--config', required=True)
    ap.add_argument('--fotos', required=True)
    ap.add_argument('--planilha')
    ap.add_argument('--template', choices=['tld', 'rc', 'vazio'], default='vazio')
    ap.add_argument('--saida', default='Relatorio_N2.docx')
    ap.add_argument('--pdf', action='store_true')
    args = ap.parse_args()

    with open(args.config, encoding='utf-8') as f:
        data = json.load(f)
    cfg = Config()
    for k, v in data.items():
        if k == 'componentes':
            cfg.componentes = []
            for c in v:
                cmp = Componente(titulo=c['titulo'], introducao=c.get('introducao',''),
                                  layout=c.get('layout', '1'))
                for fd in c.get('fotos', []):
                    cmp.fotos.append(Foto(**fd))
                cfg.componentes.append(cmp)
        elif k == 'fotos_panoramicas':
            cfg.fotos_panoramicas = [Foto(**ff) for ff in v]
        elif hasattr(cfg, k):
            setattr(cfg, k, v)

    if args.template == 'tld' and not cfg.componentes:
        cfg.componentes = [Componente(titulo=t, layout="1") for t in COMPONENTES_TLD]
    elif args.template == 'rc' and not cfg.componentes:
        cfg.componentes = [Componente(titulo=t, layout="1") for t in COMPONENTES_RC]

    fotos_dir = Path(args.fotos)
    if args.planilha:
        n = importar_fotos_csv(Path(args.planilha), cfg.componentes)
        print(f"[OK] {n} fotos via planilha.")
    else:
        n = importar_fotos_por_pasta(fotos_dir, cfg.componentes)
        print(f"[OK] {n} fotos automaticas.")

    builder = N2ReportBuilder(cfg, fotos_dir)
    builder.build()
    builder.doc.save(args.saida)
    print(f"[OK] DOCX gerado: {args.saida}")
    if args.pdf:
        pdf = converter_pdf(args.saida)
        if pdf:
            print(f"[OK] PDF gerado: {pdf}")


if __name__ == '__main__':
    main()
